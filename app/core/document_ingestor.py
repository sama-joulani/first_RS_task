from __future__ import annotations

import logging
import re
import uuid
from dataclasses import dataclass, field
from pathlib import Path

from qdrant_client import QdrantClient
from qdrant_client.models import Distance, PointStruct, VectorParams

from app.core.embedding_service import EmbeddingService
from app.models.documents import ChunkMetadata, DocumentMetadata, DocumentStatus
from app.utils.pdf_parser import PDFParser, ParsedDocument
from app.utils.text_chunker import Chunk, TextChunker

logger = logging.getLogger(__name__)


@dataclass
class IngestResult:
    document_id: str
    chunk_count: int
    page_count: int
    status: DocumentStatus
    errors: list[str] = field(default_factory=list)


class DocumentIngestor:
    """Orchestrates document parsing, chunking, embedding, and storage in Qdrant."""

    def __init__(
        self,
        embedding_service: EmbeddingService,
        qdrant_client: QdrantClient,
        collection_name: str,
        pdf_parser: PDFParser | None = None,
        text_chunker: TextChunker | None = None,
    ):
        self._embedding = embedding_service
        self._qdrant = qdrant_client
        self._collection = collection_name
        self._pdf_parser = pdf_parser or PDFParser()
        self._chunker = text_chunker or TextChunker()
        self._ensure_collection()

    def ingest(self, file_path: str | Path, metadata: DocumentMetadata) -> IngestResult:
        file_path = Path(file_path)
        errors: list[str] = []

        # Extract file metadata
        metadata.file_size = file_path.stat().st_size
        metadata.file_type = file_path.suffix.lower()

        parsed = self._parse_file(file_path)
        if not parsed or not parsed.pages:
            return IngestResult(
                document_id=metadata.document_id,
                chunk_count=0,
                page_count=0,
                status=DocumentStatus.FAILED,
                errors=["No content extracted from document"],
            )

        # Extract metadata from PDF
        if not metadata.title and parsed.metadata.get("title"):
            metadata.title = parsed.metadata["title"]
        if not metadata.author and parsed.metadata.get("author"):
            metadata.author = parsed.metadata["author"]
        if not metadata.subject and parsed.metadata.get("subject"):
            metadata.subject = parsed.metadata["subject"]
        metadata.page_count = parsed.page_count

        # Detect language from first page
        if parsed.pages:
            metadata.language = self._detect_language(parsed.pages[0].text)

        pages_data = [
            {"text": p.text, "page": p.page_number, "is_ocr": p.is_ocr}
            for p in parsed.pages
        ]
        chunks = self._chunker.chunk_pages(
            [{"text": p["text"], "page_number": p["page"]} for p in pages_data]
        )

        # Detect sections from chunk text
        sections = self._detect_sections(chunks)

        for i, chunk in enumerate(chunks):
            section = sections.get(i, "")
            # Extract tags from chunk text
            chunk_tags = self._extract_tags(chunk.text, metadata.title)
            # Merge with existing tags
            all_tags = list(set(metadata.tags + chunk_tags))
            chunk.metadata.update(
                {
                    "document_id": metadata.document_id,
                    "title": metadata.title,
                    "author": metadata.author,
                    "path": str(file_path),
                    "tags": all_tags,
                    "url": metadata.url,
                    "section": section,
                    "page": chunk.metadata.get("page_number", 0),
                    "category": metadata.category,
                    "language": metadata.language,
                }
            )

        chunk_count = self._embed_and_store(chunks)

        return IngestResult(
            document_id=metadata.document_id,
            chunk_count=chunk_count,
            page_count=parsed.page_count,
            status=DocumentStatus.INDEXED,
        )

    def _detect_language(self, text: str) -> str:
        """Detect document language."""
        # Simple heuristic: check for Arabic characters
        arabic_chars = sum(1 for c in text if '\u0600' <= c <= '\u06FF')
        total_chars = len(text.strip())
        if total_chars > 0 and arabic_chars / total_chars > 0.3:
            return "ar"
        return "en"

    def _extract_tags(self, text: str, title: str) -> list[str]:
        """Extract relevant tags from text content."""
        text_lower = text.lower()
        tags = set()
        
        # Company-related tags
        company_keywords = {
            "realsoft": ["company", "realsoft", "organization"],
            "mission": ["mission", "vision", "values", "goals"],
            "about": ["about", "company info", "overview"],
        }
        
        # Product tags
        product_keywords = {
            "al-khwarizmi": ["al-khwarizmi", "khwarizmi", "الخوارزمي"],
            "falconmap": ["falconmap", "falcon map", "falcon"],
            "realdata": ["realdata", "real data", "data hub", "data flow"],
            "adaa": ["adaa", "ada'a", "performance"],
        }
        
        # Service tags
        service_keywords = {
            "statistics": ["statistical", "statistics", "census", "survey"],
            "gis": ["gis", "mapping", "geographic", "spatial"],
            "digital": ["digital transformation", "digitization", "automation"],
            "consulting": ["consulting", "consultant", "advisory"],
            "outsourcing": ["outsourcing", "talent", "resources"],
        }
        
        # Partner tags
        partner_keywords = {
            "microsoft": ["microsoft", "azure", "cloud"],
            "esri": ["esri", "arcgis"],
            "mendix": ["mendix", "low-code", "low code"],
        }
        
        # Check each category
        all_keywords = {
            **{k: v for k, v in company_keywords.items()},
            **{k: v for k, v in product_keywords.items()},
            **{k: v for k, v in service_keywords.items()},
            **{k: v for k, v in partner_keywords.items()},
        }
        
        for tag, keywords in all_keywords.items():
            if any(kw in text_lower for kw in keywords):
                tags.add(tag)
        
        # Add title-based tags
        title_lower = title.lower()
        if "flyer" in title_lower or "brochure" in title_lower:
            tags.add("marketing")
        if "content" in title_lower:
            tags.add("content")
        
        # Add language tag
        if self._detect_language(text) == "ar":
            tags.add("arabic")
        else:
            tags.add("english")
        
        return sorted(list(tags))

    def _detect_sections(self, chunks: list[Chunk]) -> dict[int, str]:
        """Detect section headers from chunk text."""
        sections: dict[int, str] = {}
        current_section = ""

        # Common section header patterns
        section_patterns = [
            r'^(?:Chapter|Section|Part)\s+\d+[:\s]+(.+)$',  # Chapter 1: Title
            r'^\d+\.\d*\s+([A-Z][^.]+)$',  # 1. Introduction
            r'^([A-Z][A-Z\s]{3,50})$',  # ALL CAPS HEADERS
            r'^(?:Introduction|Conclusion|Summary|References|Appendix)$',  # Common sections
        ]

        for i, chunk in enumerate(chunks):
            text = chunk.text.strip()
            first_line = text.split('\n')[0] if text else ""

            for pattern in section_patterns:
                match = re.match(pattern, first_line, re.IGNORECASE)
                if match:
                    current_section = match.group(1) if match.groups() else first_line
                    break

            sections[i] = current_section

        return sections

    def delete_document(self, document_id: str) -> bool:
        from qdrant_client.models import FieldCondition, Filter, MatchValue

        try:
            self._qdrant.delete(
                collection_name=self._collection,
                points_selector=Filter(
                    must=[FieldCondition(key="document_id", match=MatchValue(value=document_id))]
                ),
            )
            return True
        except Exception:
            logger.exception("Failed to delete document %s from Qdrant", document_id)
            return False

    def _parse_file(self, file_path: Path) -> ParsedDocument | None:
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return self._pdf_parser.parse(file_path)
        elif suffix == ".docx":
            return self._parse_docx(file_path)
        else:
            logger.warning("Unsupported file type: %s", suffix)
            return None

    def _parse_docx(self, file_path: Path) -> ParsedDocument:
        from docx import Document as DocxDocument
        from app.utils.pdf_parser import PageContent

        doc = DocxDocument(str(file_path))
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        props = doc.core_properties
        metadata = {
            "title": props.title or "",
            "author": props.author or "",
        }
        pages = [PageContent(page_number=1, text=full_text)] if full_text else []
        return ParsedDocument(pages=pages, metadata=metadata, page_count=1)

    def _embed_and_store(self, chunks: list[Chunk]) -> int:
        if not chunks:
            return 0

        texts = [c.text for c in chunks]
        vectors = self._embedding.embed_batch(texts)

        points = [
            PointStruct(
                id=str(uuid.uuid4()),
                vector=vec,
                payload={
                    "text": chunk.text,
                    "chunk_index": chunk.chunk_index,
                    **{k: v for k, v in chunk.metadata.items() if k != "page_number"},
                },
            )
            for chunk, vec in zip(chunks, vectors)
        ]

        batch_size = 100
        for i in range(0, len(points), batch_size):
            self._qdrant.upsert(
                collection_name=self._collection,
                points=points[i : i + batch_size],
            )

        return len(points)

    def _ensure_collection(self) -> None:
        collections = [c.name for c in self._qdrant.get_collections().collections]
        if self._collection not in collections:
            self._qdrant.create_collection(
                collection_name=self._collection,
                vectors_config=VectorParams(
                    size=self._embedding.dimension,
                    distance=Distance.COSINE,
                ),
            )
            logger.info("Created Qdrant collection: %s", self._collection)
